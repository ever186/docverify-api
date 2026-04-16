"""
app.py — Verificador de coherencia .docx
"""
import os, re, io, traceback, logging
logging.basicConfig(level=logging.INFO)
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from spellchecker import SpellChecker

app = Flask(__name__)
CORS(app, origins=["https://ever186.github.io"])
app.config['MAX_CONTENT_LENGTH'] = 15 * 1024 * 1024

# Pre-cargar spellchecker una sola vez al arrancar
try:
    SPELL = SpellChecker(language='es')
except Exception:
    SPELL = None

ALLOWED = {'docx'}

# CRÍTICO: flask-CORS no agrega headers en errores 500
# Este after_request los fuerza en TODA respuesta incluyendo crashes
@app.after_request
def cors_always(response):
    origin = request.headers.get('Origin', '')
    response.headers['Access-Control-Allow-Origin']  = origin or '*'
    response.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    return response

# Capturar cualquier excepción no manejada y devolverla como JSON con CORS
@app.errorhandler(500)
def internal_error(e):
    resp = jsonify({'error': f'Error interno del servidor: {str(e)}'})
    resp.status_code = 500
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

@app.errorhandler(Exception)
def unhandled(e):
    resp = jsonify({'error': f'Error inesperado: {str(e)}'})
    resp.status_code = 500
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp
def allowed_file(f): return '.' in f and f.rsplit('.',1)[1].lower() in ALLOWED

# ── HELPERS ──────────────────────────────────────────────────────────────
PLACEHOLDER = re.compile(r'^<[^>]+>$')
def es_ph(v): return bool(PLACEHOLDER.match(v.strip())) if v.strip() else False
def norm(t):  return t.lower().replace('\u2013','-').replace('\u2014','-').strip()

PATRONES_FECHA = [
    r'\b\d{1,2}/\d{1,2}/\d{4}\b',
    r'\b\d{4}/\d{2}/\d{2}\b',
    r'\b\d{1,2}-\d{1,2}-\d{4}\b',
    r'\b\d{4}-\d{2}-\d{2}\b',
    r'\b\d{1,2}\s+de\s+\w+\s+de\s+\d{4}\b',
]
def extraer_fechas(t):
    r = []
    for p in PATRONES_FECHA:
        r += re.findall(p, t, re.IGNORECASE)
    return list(set(f for f in r if not re.match(r'^\d\.\d$', f)))

def idx_conc(tablas):
    for i,t in enumerate(tablas):
        for f in t:
            for c in f:
                if 'id azure' in c.lower() or c.lower() == 'consecutivo':
                    return i
    return -1

# ── EXTRACCIÓN ────────────────────────────────────────────────────────────
def extraer_contenido(fb):
    doc   = Document(io.BytesIO(fb))
    props = doc.core_properties
    W     = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def texto_tc(tc):
        return ''.join(t.text or '' for t in tc.iter(f'{{{W}}}t')).strip()

    def leer_fila(row):
        celdas = []
        for child in row._tr:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'tc':
                celdas.append(texto_tc(child))
            elif tag == 'sdt':
                sc = child.find(f'.//{{{W}}}sdtContent')
                if sc is not None:
                    tc = sc.find(f'.//{{{W}}}tc')
                    if tc is not None:
                        celdas.append(texto_tc(tc))
                        continue
                celdas.append(''.join(t.text or '' for t in child.iter(f'{{{W}}}t')).strip())
        # fallback deduplicado si SDT no dio más columnas
        seen, uniq = [], []
        for cell in row.cells:
            if cell not in seen:
                seen.append(cell)
                uniq.append(cell.text.strip())
        return celdas if len(celdas) >= len(uniq) else uniq

    encabezados, vistos = [], set()
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            t = p.text.strip()
            if t and t not in vistos:
                encabezados.append(t); vistos.add(t)

    parrafos = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tablas   = [[leer_fila(r) for r in t.rows] for t in doc.tables]
    texto    = '\n'.join(encabezados + parrafos + [c for t in tablas for f in t for c in f])

    return {
        'meta':        {'autor': (props.author or '').strip(), 'mod_por': (props.last_modified_by or '').strip()},
        'encabezados': encabezados,
        'parrafos':    parrafos,
        'tablas':      tablas,
        'texto':       texto,
    }

# ── SECCIÓN 1: ENCABEZADO ─────────────────────────────────────────────────
def s_encabezado(c, titulo):
    vals, enc = [], c['encabezados']
    if titulo:
        ok = any(norm(titulo) in norm(e) for e in enc)
        vals.append({'estado':'OK' if ok else 'ERROR',
                     'detalle': f'Título "{titulo}" {"✔ presente" if ok else "✘ NO encontrado"} en encabezado'})
    fechas = [f for e in enc for f in extraer_fechas(e)]
    if fechas:
        vals.append({'estado':'INFO','detalle':f'Fecha en encabezado: {list(set(fechas))}'})
    else:
        vals.append({'estado':'WARN','detalle':'Sin fecha concreta en encabezado'})
    estado = 'ERROR' if any(v['estado']=='ERROR' for v in vals) else \
             'WARN'  if any(v['estado']=='WARN'  for v in vals) else 'OK'
    return {'titulo':'Encabezado','estado':estado,'fragmento':enc,'validaciones':vals,'tipo':'encabezado'}

# ── SECCIÓN 2: INFO DOCUMENTO ─────────────────────────────────────────────
def s_info(c, consecutivo):
    tablas, meta = c['tablas'], c['meta']
    vals, frag = [], []
    try:
        t1 = tablas[1]
    except:
        return {'titulo':'Información del documento','estado':'WARN','fragmento':[],
                'validaciones':[{'estado':'WARN','detalle':'Tabla de información no encontrada'}],'tipo':'tabla_info'}

    campos = {}
    for fila in t1:
        campo = fila[0].rstrip(':').strip() if fila else ''
        valor = fila[1].strip() if len(fila) > 1 else ''
        frag.append({'campo':campo,'valor':valor or '—','es_placeholder':es_ph(valor)})
        campos[campo.lower()] = valor

    if consecutivo:
        cod = campos.get('código', campos.get('codigo', ''))
        if norm(consecutivo) in norm(cod):
            vals.append({'estado':'OK','detalle':f'Consecutivo "{consecutivo}" ✔ coincide con Código'})
        elif es_ph(cod):
            vals.append({'estado':'WARN','detalle':f'Código es placeholder: "{cod}"'})
        else:
            vals.append({'estado':'ERROR','detalle':f'Consecutivo "{consecutivo}" ✘ no coincide — Código: "{cod}"'})

    autor = campos.get('autor','')
    if autor and not es_ph(autor):
        vals.append({'estado':'INFO','detalle':f'Autor metadatos: "{meta["autor"]}" | Mod. por: "{meta["mod_por"]}"'})
        coincide = meta['autor'] and (meta['autor'].lower() in autor.lower() or autor.lower() in meta['autor'].lower())
        vals.append({'estado':'OK' if coincide else 'WARN',
                     'detalle': f'Autor "{autor}" {"✔ coincide con" if coincide else "≠"} metadato "{meta["autor"]}"'})
    else:
        vals.append({'estado':'WARN','detalle':'Campo Autor es placeholder o está vacío'})

    estado = 'ERROR' if any(v['estado']=='ERROR' for v in vals) else \
             'WARN'  if any(v['estado']=='WARN'  for v in vals) else 'OK'
    return {'titulo':'Información del documento','estado':estado,'fragmento':frag,'validaciones':vals,'tipo':'tabla_info'}

# ── SECCIÓN 3: HISTORIAL ──────────────────────────────────────────────────
def s_historial(c, titulo):
    tablas = c['tablas']
    vals   = []
    try:
        th = tablas[2]
    except:
        return {'titulo':'Historial de revisiones','estado':'WARN','fragmento':{'encabezados':[],'filas':[]},
                'validaciones':[{'estado':'WARN','detalle':'Historial no encontrado'}],'tipo':'tabla_historial'}

    enc_h  = th[0] if th else []
    filas  = th[1:] if len(th) > 1 else []
    frag   = {'encabezados': enc_h, 'filas': filas}

    if titulo:
        en_hist = any(norm(titulo) in norm(c2) for f in filas for c2 in f)
        vals.append({'estado':'OK' if en_hist else 'ERROR',
                     'detalle': f'Título "{titulo}" {"✔ presente" if en_hist else "✘ NO encontrado"} en historial'})

    fechas = list(set(f for fila in filas for cel in fila if not es_ph(cel) for f in extraer_fechas(cel)))
    if fechas:
        vals.append({'estado':'OK' if len(fechas)==1 else 'ERROR',
                     'detalle': f'Fechas en historial: {fechas}' if len(fechas)>1 else f'Fecha historial: "{fechas[0]}"'})
    else:
        vals.append({'estado':'WARN','detalle':'Sin fechas concretas en historial'})

    autores = list(set(f[-1] for f in filas if f and f[-1] and not es_ph(f[-1]) and len(f[-1])>3))
    if len(autores)==1:
        vals.append({'estado':'OK','detalle':f'Autor coherente: "{autores[0]}"'})
    elif len(autores)>1:
        vals.append({'estado':'ERROR','detalle':f'Autores inconsistentes: {autores}'})
    else:
        vals.append({'estado':'WARN','detalle':'Autores en historial son placeholders'})

    estado = 'ERROR' if any(v['estado']=='ERROR' for v in vals) else \
             'WARN'  if any(v['estado']=='WARN'  for v in vals) else 'OK'
    return {'titulo':'Historial de revisiones','estado':estado,'fragmento':frag,'validaciones':vals,'tipo':'tabla_historial'}

# ── SECCIÓN 4: CONCLUSIONES ───────────────────────────────────────────────
def s_conclusiones(c, titulo, id_tarea, consecutivo):
    tablas, parrafos = c['tablas'], c['parrafos']
    vals = []
    ti   = idx_conc(tablas)

    parrafo_c = next((p for p in parrafos if 'corresponden a la aplicación' in p.lower() or 'resultados obtenidos' in p.lower()), '')
    tabla_c   = tablas[ti] if ti >= 0 else []

    frag = {'parrafo': parrafo_c, 'tabla_campos': []}
    vals_conc = {}
    for fila in (tabla_c[1:] if tabla_c else []):
        if len(fila) >= 2:
            campo = fila[0].strip()
            valor = fila[1].strip()
            frag['tabla_campos'].append({'campo':campo,'valor':valor or '—','es_placeholder':es_ph(valor)})
            vals_conc[campo.lower()] = valor

    if titulo:
        en_c = any(norm(titulo) in norm(p) for p in parrafos) or \
               any(norm(titulo) in norm(cel) for f in tabla_c for cel in f)
        vals.append({'estado':'OK' if en_c else 'WARN',
                     'detalle': f'Título "{titulo}" {"✔ en Conclusiones" if en_c else "✘ NO en Conclusiones"}'})

    if id_tarea:
        v = vals_conc.get('id azure','')
        if norm(id_tarea) in norm(v):
            vals.append({'estado':'OK','detalle':f'ID Azure "{id_tarea}" ✔ coincide'})
        elif es_ph(v):
            vals.append({'estado':'WARN','detalle':f'ID Azure es placeholder: "{v}"'})
        else:
            vals.append({'estado':'ERROR','detalle':f'ID Azure "{id_tarea}" ✘ no coincide — contiene: "{v}"'})

    if consecutivo:
        v = vals_conc.get('consecutivo','')
        if norm(consecutivo) in norm(v):
            vals.append({'estado':'OK','detalle':f'Consecutivo "{consecutivo}" ✔ en conclusiones'})
        elif es_ph(v):
            vals.append({'estado':'WARN','detalle':f'Consecutivo es placeholder: "{v}"'})
        else:
            vals.append({'estado':'ERROR','detalle':f'Consecutivo "{consecutivo}" ✘ no coincide — contiene: "{v}"'})

    resultado = vals_conc.get('resultado','')
    if resultado and not es_ph(resultado):
        vals.append({'estado':'OK' if resultado.upper()=='CERTIFICADA' else 'WARN',
                     'detalle':f'Resultado: "{resultado}"'})

    estado = 'ERROR' if any(v['estado']=='ERROR' for v in vals) else \
             'WARN'  if any(v['estado']=='WARN'  for v in vals) else 'OK'
    return {'titulo':'Conclusiones','estado':estado,'fragmento':frag,'validaciones':vals,'tipo':'conclusiones'}

# ── SECCIÓN 5: FECHAS ─────────────────────────────────────────────────────
def s_fechas(c):
    tablas, enc, parrafos = c['tablas'], c['encabezados'], c['parrafos']
    hallazgos = {}

    for e in enc:
        for f in extraer_fechas(e):
            hallazgos[f'Encabezado — {e[:50]}'] = f

    try:
        for fi in [1,2]:
            for cel in tablas[2][fi]:
                if not es_ph(cel):
                    for f in extraer_fechas(cel):
                        hallazgos[f'Historial fila {fi} — {cel[:40]}'] = f
    except: pass

    for ti,tabla in enumerate(tablas):
        if ti==2: continue
        for fi,fila in enumerate(tabla):
            for cel in fila:
                if not es_ph(cel):
                    for f in extraer_fechas(cel):
                        hallazgos[f'Tabla {ti} fila {fi} — {cel[:30]}'] = f

    for i,p in enumerate(parrafos):
        for f in extraer_fechas(p):
            hallazgos[f'Párrafo — {p[:50]}'] = f

    if not hallazgos:
        return {'titulo':'Coherencia de fechas','estado':'WARN',
                'fragmento':[],'validaciones':[{'estado':'WARN','detalle':'Sin fechas concretas — pueden ser placeholders'}],
                'tipo':'fechas'}

    frag  = [{'ubicacion':u,'fecha':f} for u,f in hallazgos.items()]
    uniq  = list(set(hallazgos.values()))
    vals  = []
    if len(uniq)==1:
        vals.append({'estado':'OK','detalle':f'Fechas coherentes: "{uniq[0]}"'})
    else:
        vals.append({'estado':'ERROR','detalle':f'Fechas inconsistentes: {uniq}'})
        for u,f in hallazgos.items():
            vals.append({'estado':'INFO','detalle':f'"{f}" en {u}'})

    estado = 'ERROR' if any(v['estado']=='ERROR' for v in vals) else 'OK'
    return {'titulo':'Coherencia de fechas','estado':estado,'fragmento':frag,'validaciones':vals,'tipo':'fechas'}

# ── SECCIÓN 6: ORTOGRAFÍA ─────────────────────────────────────────────────
IGNORAR = {
    'sanitizar','canonicalización','csrf','xss','sql','api','backend','frontend',
    'payload','bypass','exploit','apikey','endpoint','endpoints','token','bearer',
    'oauth','jwt','https','http','cors','json','xml','rest','soap','headers',
    'injection','encoding','hashing','hash','access','application','mobile',
    'cloud','devops','aplicaciones','autenticación','autorización','controles',
    'funcionalidades','vulnerabilidades','caracteres','conclusiones','autorizados',
    'apropiados','corresponden','ficheros','evaluada','incluyendo','periódica',
    'telecomunicaciones','informática','módulos','distribución','subversión',
    'versión','sección','también','través','según','opciones','deberán',
    'tendrán','están','acceso','usuarios','servidores','sistemas','seguridad',
    'informe','documentación','actualización','revisión','creación','generación',
    'aprobación','modificación','versiones','registros','información','gerencia',
    'dirección','recomendaciones','implementación','validación','implementando',
    'modificaciones','observaciones','obtenidos','realizadas','relacionadas',
    'revisiones','necesarios','privilegios','recursos','sesiones','siguientes',
    'áreas','adicionalmente','accesibles','ambientes','credenciales','canales',
    'clientes','activación','activos','aplicacion','aceptando',
}

def es_hex(p): return bool(re.match(r'^[0-9a-f]{4,}$', p))

def s_ortografia(c):
    if not SPELL:
        return {'titulo':'Ortografía','estado':'WARN','fragmento':[],
                'validaciones':[{'estado':'WARN','detalle':'Revisor ortográfico no disponible'}],'tipo':'ortografia'}

    texto = re.sub(r'<[^>]+>',' ',c['texto'])
    texto = re.sub(r'https?://\S+',' ',texto)
    texto = re.sub(r'\b[A-Z]{2,}\b',' ',texto)
    texto = re.sub(r'\b\d[\d.,/-]*\b',' ',texto)
    texto = re.sub(r'[^\w\sáéíóúÁÉÍÓÚñÑüÜ]',' ',texto)
    texto = re.sub(r'\s+',' ',texto).strip()

    palabras = [p.lower() for p in texto.split() if len(p) > 3]
    revisar  = [p for p in set(palabras) if p not in IGNORAR and not es_hex(p)][:150]

    if not revisar:
        return {'titulo':'Ortografía','estado':'OK','fragmento':[],
                'validaciones':[{'estado':'OK','detalle':'Sin errores ortográficos'}],'tipo':'ortografia'}

    try:
        desc = SPELL.unknown(revisar)
        errores = [p for p in desc if p not in IGNORAR and not es_hex(p)]
        sugs = []
        for p in sorted(errores)[:25]:
            s = SPELL.correction(p)
            sugs.append({'palabra':p,'sugerencia': s if s and s!=p else None})
    except Exception as e:
        return {'titulo':'Ortografía','estado':'WARN','fragmento':[],
                'validaciones':[{'estado':'WARN','detalle':f'Error ortografía: {str(e)}'}],'tipo':'ortografia'}

    vals = [{'estado':'OK' if not sugs else 'INFO',
             'detalle':'Sin errores ortográficos' if not sugs else f'{len(sugs)} posible(s) error(es) — solo sugerencias'}]
    return {'titulo':'Ortografía','estado':'OK' if not sugs else 'INFO',
            'fragmento':sugs,'validaciones':vals,'tipo':'ortografia'}

# ── ENDPOINT ──────────────────────────────────────────────────────────────
@app.route('/verificar', methods=['POST','OPTIONS'])
def verificar():
    if request.method == 'OPTIONS':
        return jsonify({}), 200

    if 'file' not in request.files:
        return jsonify({'error':'No se recibió ningún archivo'}), 400

    file = request.files['file']
    if not file or not file.filename or not allowed_file(file.filename):
        return jsonify({'error':'Solo se aceptan archivos .docx'}), 400

    fb = file.read()
    if not fb:
        return jsonify({'error':'El archivo está vacío'}), 400

    titulo_p    = request.form.get('titulo','').strip()
    id_tarea    = request.form.get('id_tarea','').strip()
    consecutivo = request.form.get('consecutivo','').strip()

    try:
        c = extraer_contenido(fb)
    except Exception as e:
        logging.error(f'extraer_contenido error: {traceback.format_exc()}')
        resp = jsonify({'error':f'No se pudo leer el documento: {str(e)}'})
        resp.status_code = 422
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp

    secciones = []
    for fn, args, nombre in [
        (s_encabezado,    (c, titulo_p),                        'Encabezado'),
        (s_info,          (c, consecutivo),                     'Información del documento'),
        (s_historial,     (c, titulo_p),                        'Historial de revisiones'),
        (s_conclusiones,  (c, titulo_p, id_tarea, consecutivo), 'Conclusiones'),
        (s_fechas,        (c,),                                  'Coherencia de fechas'),
        (s_ortografia,    (c,),                                  'Ortografía'),
    ]:
        try:
            secciones.append(fn(*args))
        except Exception as e:
            logging.error(f'Sección {nombre} error: {traceback.format_exc()}')
            secciones.append({'titulo':nombre,'estado':'WARN','fragmento':[],
                              'validaciones':[{'estado':'WARN','detalle':f'Error: {str(e)}'}],'tipo':'error'})

    ok   = sum(1 for s in secciones if s['estado']=='OK')
    warn = sum(1 for s in secciones if s['estado'] in ('WARN','INFO'))
    err  = sum(1 for s in secciones if s['estado']=='ERROR')

    return jsonify({'secciones':secciones,'resumen':{'ok':ok,'alertas':warn,'errores':err}})

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status':'ok'})

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT',5000)), debug=False)
